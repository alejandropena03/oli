from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Consultor Estrategico Codex"
ASSET_DIR = OUT_DIR / "assets_10_sales_ux"
DOCX_PATH = OUT_DIR / "10_experiencia_usuario_oli_sales_automation_team.docx"


def add_title(doc: Document, text: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(18, 24, 38)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(86, 98, 118)


def add_h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(18, 24, 38)


def add_p(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def save_flow_diagram(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13.33, 7.5), dpi=150)
    ax.set_facecolor("#f7f8fb")
    fig.patch.set_facecolor("#f7f8fb")
    ax.axis("off")

    nodes = [
        ("Wake word\nOpenWakeWord", 0.05, 0.72, "#d7f5ee"),
        ("VAD + Whisper\nvoz -> texto", 0.22, 0.72, "#d7f5ee"),
        ("Mission Kernel\nestado + permisos", 0.39, 0.72, "#e6ecff"),
        ("Model Router\nTeam H100 + APIs", 0.56, 0.72, "#e6ecff"),
        ("Cloud Runtime\nworkspace remoto", 0.73, 0.72, "#fff2cc"),
        ("Tool Router\nAPIs / browser / shell", 0.39, 0.42, "#eaf7d5"),
        ("Attio CRM", 0.56, 0.42, "#ffffff"),
        ("Apollo Data", 0.70, 0.42, "#ffffff"),
        ("Gmail Workspace", 0.84, 0.42, "#ffffff"),
        ("Smartlead\nsecuencias opcionales", 0.70, 0.18, "#ffffff"),
        ("Evidence + Audit\nvalidacion + replay", 0.39, 0.18, "#f4dcff"),
        ("Usuario\naprobaciones", 0.12, 0.18, "#ffd9d9"),
    ]

    for label, x, y, color in nodes:
        box = plt.Rectangle((x, y), 0.13, 0.12, linewidth=1.2, edgecolor="#273044", facecolor=color)
        ax.add_patch(box)
        ax.text(x + 0.065, y + 0.06, label, ha="center", va="center", fontsize=9, color="#182033")

    arrows = [
        ((0.18, 0.78), (0.22, 0.78)),
        ((0.35, 0.78), (0.39, 0.78)),
        ((0.52, 0.78), (0.56, 0.78)),
        ((0.69, 0.78), (0.73, 0.78)),
        ((0.455, 0.72), (0.455, 0.54)),
        ((0.52, 0.48), (0.56, 0.48)),
        ((0.69, 0.48), (0.70, 0.48)),
        ((0.83, 0.48), (0.84, 0.48)),
        ((0.765, 0.42), (0.765, 0.30)),
        ((0.455, 0.42), (0.455, 0.30)),
        ((0.39, 0.24), (0.25, 0.24)),
        ((0.18, 0.24), (0.12, 0.24)),
    ]
    for start, end in arrows:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=1.2, color="#273044"))

    ax.text(0.05, 0.93, "Experiencia final Oli: solicitud minima -> sistema B2B operativo", fontsize=18, weight="bold", color="#121826")
    ax.text(0.05, 0.89, "Default corregido: cloud-managed, Attio + Apollo, aprobaciones humanas para comunicaciones externas.", fontsize=10, color="#566276")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def save_pipeline_diagram(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13.33, 7.5), dpi=150)
    ax.set_facecolor("#fbfaf7")
    fig.patch.set_facecolor("#fbfaf7")
    ax.axis("off")

    stages = [
        ("1\nDescubrir ICP", "Oli infiere desde dominio Gmail,\nweb publica y memoria vacia"),
        ("2\nApollo", "Busca cuentas y personas\npor industria, seniority y pais"),
        ("3\nEnriquecer", "Apollo enrichment + web evidence\nsin inventar emails"),
        ("4\nScoring", "Modelo OSS en GPU Team\ncalifica fit, urgencia y riesgo"),
        ("5\nAttio", "Crea objetos, listas,\nstages y propiedades"),
        ("6\nOutreach", "Gmail draft o Smartlead\ncon aprobacion humana"),
        ("7\nAprender", "Replies -> clasificacion ->\nmemoria + playbook"),
    ]

    x0 = 0.05
    w = 0.12
    gap = 0.015
    for idx, (title, desc) in enumerate(stages):
        x = x0 + idx * (w + gap)
        y = 0.50
        color = "#e6ecff" if idx in (0, 3, 6) else "#ffffff"
        ax.add_patch(plt.Rectangle((x, y), w, 0.22, facecolor=color, edgecolor="#273044", linewidth=1.1))
        ax.text(x + w / 2, y + 0.16, title, ha="center", va="center", fontsize=11, weight="bold", color="#121826")
        ax.text(x + w / 2, y + 0.065, desc, ha="center", va="center", fontsize=7.5, color="#465066")
        if idx < len(stages) - 1:
            ax.annotate("", xy=(x + w + gap * 0.8, y + 0.11), xytext=(x + w, y + 0.11), arrowprops=dict(arrowstyle="->", lw=1.2))

    ax.text(0.05, 0.88, "Pipeline operativo que Oli construiria", fontsize=18, weight="bold", color="#121826")
    ax.text(0.05, 0.84, "El sistema no empieza por enviar emails. Empieza por crear una maquina auditable de datos, CRM, scoring y aprobacion.", fontsize=10, color="#566276")
    ax.text(0.05, 0.28, "Regla de seguridad: ninguna comunicacion externa sale sin aprobacion; los primeros runs crean borradores, tests y evidencia.", fontsize=10, color="#8a3b12")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    flow_img = ASSET_DIR / "workflow_oli_sales.png"
    pipeline_img = ASSET_DIR / "pipeline_sales.png"
    save_flow_diagram(flow_img)
    save_pipeline_diagram(pipeline_img)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)

    add_title(
        doc,
        "Experiencia simulada: Oli crea un sistema de automatizacion de ventas B2B",
        "Paquete Team | Usuario sin stack previo | Gmail empresarial conectado | Vision final segun TDD + correccion cloud-managed",
    )

    add_h(doc, "0. Correccion de criterio", 1)
    add_p(
        doc,
        "El stack por defecto no debe ser HubSpot + n8n. Ese stack solo aparece si el usuario ya lo tiene o si hay una razon operacional fuerte. "
        "Para un usuario Team que no tiene nada montado, Oli debe proponer una arquitectura moderna, auditable y API-first: Attio como CRM flexible, "
        "Apollo como base de datos/prospeccion, Gmail empresarial como identidad inicial, Smartlead o Instantly para cold outbound solo cuando el usuario aprueba "
        "infraestructura de deliverability, y Oli Cloud Runtime como operador que conecta, construye, prueba y aprende."
    )
    add_p(
        doc,
        "Esto es una simulacion de la experiencia final, no de la V0 actual del repo. La V0 actual solo implementa una base del Mission Kernel; "
        "este documento muestra como deberia sentirse Oli cuando el TDD este construido completo."
    )

    doc.add_picture(str(flow_img), width=Inches(7.2))

    add_h(doc, "1. Escenario exacto", 1)
    add_table(
        doc,
        ["Variable", "Supuesto de la simulacion"],
        [
            ["Plan", "Team: 4+ seats, memoria de empresa, roles, audit exportable, mayor concurrencia, GPU H100 segun pricing-model-v6."],
            ["Usuario", "No tecnico. Solo dice: 'Oli, creame el sistema de automatizacion de ventas B2B'."],
            ["Stack inicial", "No tiene CRM, no tiene Apollo conectado, no tiene Smartlead, no tiene Slack, no tiene landing ni base de leads."],
            ["Lo que si tiene", "Gmail empresarial / Google Workspace y creditos basicos para Claude Code."],
            ["Principio de Oli", "Investigar antes de preguntar; preguntar maximo lo critico; no enviar comunicaciones externas sin aprobacion."],
        ],
    )

    add_h(doc, "2. Lo que el usuario vive", 1)
    add_p(doc, "00:00 - Wake word")
    add_bullets(
        doc,
        [
            "El usuario dice: 'Oli, creame el sistema de automatizacion de ventas B2B'.",
            "La app local ligera escucha con OpenWakeWord. El audio no viaja todavia; solo detecta la palabra de activacion.",
            "VAD corta la frase completa. Whisper transcribe: 'creame el sistema de automatizacion de ventas B2B'.",
            "La mision aparece en Mission Control: intake_received -> interpreting_intent.",
        ],
    )

    add_p(doc, "00:05 - Oli no pregunta de entrada por CRM")
    add_p(
        doc,
        "Como el usuario no tiene nada montado, Oli no pregunta 'que CRM usas'. Detecta que la respuesta es 'ninguno' por contexto de onboarding. "
        "La pregunta correcta seria otra: si puede crear el stack recomendado. Pero antes investiga el dominio de Gmail, la web de la empresa si existe, "
        "y la memoria de empresa. Si no hay datos, crea un sistema base con supuestos editables."
    )

    add_p(doc, "00:20 - Ruta propuesta visible")
    add_table(
        doc,
        ["Modulo", "Decision que Oli propone", "Por que"],
        [
            ["CRM", "Attio", "CRM flexible, API-first, objetos/listas custom, buena base para equipos pequenos y workflows modernos."],
            ["Base de prospectos", "Apollo", "Busqueda de personas y organizaciones, enrichment, contactos, cuentas y secuencias via API segun plan."],
            ["Outbound", "Gmail drafts al inicio; Smartlead/Instantly si aprueban cold outbound real", "No quemar el dominio principal. Primero borradores y validacion; luego infraestructura de envio."],
            ["Scoring", "Oli scoring engine sobre GPU Team + evidencias web", "No depender solo de campos de Apollo. Oli razona fit, urgencia, sector y trigger."],
            ["Automatizacion", "Oli Playbook + workers propios; no n8n default", "n8n es opcion si existe. El default debe ser producto Oli, no externalizar la experiencia."],
            ["Evidencia", "Mission Replay + audit trail", "Cada lead, score, draft, API call y aprobacion queda trazable."],
        ],
    )

    add_h(doc, "3. Lo que ocurre por debajo", 1)
    doc.add_picture(str(pipeline_img), width=Inches(7.2))

    add_p(doc, "Fase A - Intake tecnico invisible")
    add_bullets(
        doc,
        [
            "Mission Kernel crea mission_id y estado intake_received.",
            "InterpretIntent usa modelo rapido: goal=build_b2b_sales_automation_system.",
            "RetrieveContext busca memoria personal, memoria de empresa, dominio de Gmail y misiones anteriores. Como esta vacio, marca cold_start=true.",
            "Permission model calcula clase maxima 3 porque eventualmente habra comunicaciones externas. Todo envio queda bloqueado hasta aprobacion.",
            "Model Router asigna GPU Team H100 para razonamiento local/open-source y Tier 3 API solo para arquitectura critica o Claude Code si hay coding complejo.",
        ],
    )

    add_p(doc, "Fase B - Onboarding de herramientas sin tokens visibles")
    add_bullets(
        doc,
        [
            "Oli abre OAuth/connection flows para Google Workspace, Attio, Apollo y opcionalmente Smartlead.",
            "El usuario ve botones: Conectar Attio, Conectar Apollo, Crear workspace Attio, Crear pipeline base.",
            "Credential Broker guarda credenciales. El LLM nunca ve tokens.",
            "Si Attio o Apollo no existen, Oli puede abrir el signup por browser automation y dejarlo listo hasta donde sea legal/permitido; si requiere pago, pide aprobacion.",
            "Si Apollo API requiere plan especifico, Oli lo marca como bloqueo comercial y propone modo degradado: Google Sheets + web research + Attio manual import.",
        ],
    )

    add_p(doc, "Fase C - Diseno del sistema")
    add_table(
        doc,
        ["Objeto en Attio", "Campos que Oli crea", "Uso"],
        [
            ["Companies", "ICP tier, pais, industria, empleados, fuente, fit_score, trigger, owner, stage", "Cuenta objetivo y contexto comercial."],
            ["People", "rol, seniority, email_status, source, persona_fit, last_touch, next_action", "Contacto individual."],
            ["Opportunities", "pipeline_stage, expected_value, confidence, next_step, meeting_status", "Proceso comercial."],
            ["Campaigns", "segmento, messaging_angle, status, approved_by, send_policy", "Gobierno de outreach."],
            ["Evidence", "source_url, captured_at, claim, confidence", "Por que Oli cree que ese lead vale."],
        ],
    )

    add_h(doc, "4. Mock data realista con empresas reales", 1)
    add_p(
        doc,
        "Para no inventar datos personales, el mock usa empresas reales y roles objetivo. Los nombres/emails personales quedan como 'por descubrir en Apollo' "
        "porque Oli no debe fabricar correos. La busqueda real vendria de Apollo People Search + Enrichment con aprobacion y creditos disponibles."
    )
    add_table(
        doc,
        ["Cuenta", "Dato publico usado", "Persona objetivo", "Por que entra al ICP"],
        [
            ["Simetrik", "SaaS de automatizacion de conciliacion y reporting financiero.", "VP Sales / RevOps / Partnerships", "Empresa B2B SaaS con venta consultiva y alto valor por pipeline."],
            ["Alegra", "Ecosistema de software contable/administrativo para pymes en LatAm.", "Head of Growth / Sales Ops", "Volumen alto de leads pymes, necesidad de segmentacion y automatizacion."],
            ["Frubana", "Marketplace B2B para restaurantes y proveedores en LatAm.", "Head of B2B Growth / Revenue", "Operaciones B2B con cuentas, mercados y expansion regional."],
            ["Liftit", "Tecnologia logistica para mover cargas y entregas en America Latina.", "Head of Sales / Enterprise Logistics", "Venta B2B a empresas medianas/grandes, ciclos de venta consultivos."],
        ],
    )

    add_h(doc, "5. Plan que Oli mostraria antes de actuar", 1)
    add_p(
        doc,
        "/oli ruta propuesta\n"
        "Objetivo: crear un sistema B2B de ventas desde cero.\n"
        "Stack recomendado: Attio CRM + Apollo prospecting + Oli scoring + Gmail drafts + Smartlead opcional.\n"
        "Datos que salen: busquedas de empresas/personas en Apollo; borradores en Gmail; registros en Attio.\n"
        "No envio emails sin aprobacion.\n"
        "Costo estimado: 45-75 minutos de compute Team; 0-20 creditos de Apollo si enriquecemos contactos; Smartlead solo si apruebas.\n"
        "Riesgo: deliverability si usamos tu dominio principal. Recomendacion: empezar con drafts y configurar dominios outbound antes de envio masivo.\n"
        "Punto de aprobacion: conectar herramientas, comprar/usar creditos, enviar cualquier comunicacion externa."
    )

    add_h(doc, "6. Ejecucion detallada", 1)
    add_table(
        doc,
        ["Paso", "Herramienta", "Accion", "Estado visible"],
        [
            ["1", "Google Workspace", "Lee dominio, firma, posibles contactos previos y calendario si se autoriza.", "Conectando contexto base."],
            ["2", "Web research", "Busca web de la empresa del usuario y extrae propuesta de valor inicial.", "Entendiendo que vendes."],
            ["3", "Attio API", "Crea workspace schema: Companies, People, Opportunities, Campaigns, Evidence.", "CRM base creado."],
            ["4", "Apollo API", "Ejecuta Organization Search y People Search por segmentos definidos.", "Buscando prospectos."],
            ["5", "Apollo Enrichment", "Enriquece solo los leads que pasan fit preliminar para no gastar creditos basura.", "Enriqueciendo candidatos."],
            ["6", "Oli GPU Team", "Scorea cuentas/contactos con ICP fit, trigger, seniority y evidencia.", "Priorizando."],
            ["7", "Attio API", "Crea cuentas/contactos/opportunities con links a evidencia.", "CRM poblado."],
            ["8", "Gmail API", "Genera borradores 1:1, no envia.", "Borradores listos para revision."],
            ["9", "Smartlead API opcional", "Si se aprueba cold outbound, crea campaign, secuencia, schedule y mailbox rotation.", "Secuencia preparada, pausada hasta aprobacion."],
            ["10", "ValidationSuboperator", "Corre tests: dedupe, campos obligatorios, links, scoring, ausencia de envio sin aprobacion.", "Validando entrega."],
        ],
    )

    add_h(doc, "7. Tests que Oli deberia correr", 1)
    add_bullets(
        doc,
        [
            "No se creo ningun contacto duplicado en Attio.",
            "Cada cuenta tiene al menos una evidencia publica o Apollo source.",
            "Ningun email fue enviado; solo se crearon drafts o campanas pausadas.",
            "Los leads con score alto tienen buyer persona plausible y empresa dentro del ICP.",
            "Los leads con datos incompletos quedan en estado 'needs_research', no pasan a outreach.",
            "El sistema respeta Gmail/Workspace y no usa el dominio principal para cold outbound masivo sin configurar deliverability.",
        ],
    )

    add_h(doc, "8. Comprobante final que veria el usuario", 1)
    add_p(
        doc,
        "/oli comprobante\n"
        "Entregado:\n"
        "- CRM Attio configurado con pipeline B2B.\n"
        "- 4 segmentos iniciales creados: SaaS LatAm, fintech ops, logistics tech, B2B marketplaces.\n"
        "- 120 cuentas candidatas encontradas en Apollo; 37 pasaron scoring; 12 listas para revision humana.\n"
        "- 12 borradores Gmail generados, ninguno enviado.\n"
        "- Campana Smartlead preparada en pausa, pendiente de infraestructura de dominio/inboxes.\n"
        "- Playbook 'b2b_sales_system_v1' propuesto para reutilizar.\n"
        "Revision humana:\n"
        "- Aprobar ICP final.\n"
        "- Aprobar o editar los primeros 12 borradores.\n"
        "- Decidir si se configura cold outbound real con dominios dedicados.\n"
        "Riesgo marcado:\n"
        "- No usar Gmail principal para volumen alto."
    )

    add_h(doc, "9. Diferencia entre stack malo y stack Oli", 1)
    add_table(
        doc,
        ["Decision floja", "Decision Oli"],
        [
            ["HubSpot porque todo el mundo lo usa", "Attio por flexibilidad API-first y fit con equipos modernos."],
            ["n8n como cerebro", "Oli como cerebro; n8n solo si el usuario ya lo tiene o conviene por mantenimiento visual."],
            ["Enviar correos de una vez", "Borradores, aprobacion, dominio/inbox strategy y ramp-up."],
            ["Comprar leads sin filtrar", "Apollo search barato primero, enrichment solo despues de fit preliminar."],
            ["Score por formula simple", "Score con evidencia, ICP, seniority, trigger y razonamiento del modelo."],
            ["Reporte bonito", "Sistema operativo: CRM, data, drafts, scoring, validation, audit, playbook."],
        ],
    )

    add_h(doc, "10. Fuentes usadas para la simulacion", 1)
    add_bullets(
        doc,
        [
            "TDD Oli: ADR-004 permission model, ADR-014 credentials, ADR-016 model routing/GPU, ADR-018 teams, ADR-020 tool security, slice-002 sales automation.",
            "Attio Docs: REST API y Developer Platform para leer/escribir workspace data y construir apps/integrations.",
            "Apollo Docs: People API Search, Organization Search, Enrichment, Contacts, Accounts, Deals, Sequences y API access por plan.",
            "Smartlead API Docs: campaign architecture con email accounts, sequences, leads y schedule; API para campaign management y analytics.",
            "Google Gmail API Docs: users.messages.send y OAuth scopes para Gmail.",
            "Fuentes publicas de empresas mock: Simetrik, Alegra, Frubana, Liftit.",
        ],
    )

    add_h(doc, "11. Veredicto de producto", 1)
    add_p(
        doc,
        "Si Oli esta en paquete Team y el usuario pide algo tan amplio como 'creame el sistema de automatizacion de ventas B2B', "
        "la respuesta correcta no es una charla ni una receta. Es una ruta ejecutable con stack recomendado, aprobaciones, onboarding de herramientas, "
        "datos reales, scoring, CRM poblado, borradores listos, pruebas y playbook. El usuario no deberia tener que saber que es Attio API, Apollo enrichment, "
        "OAuth PKCE, H100, E2B o Mission Kernel. Solo deberia ver que Oli convirtio una frase vaga en una maquina comercial inicial, segura y auditable."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
